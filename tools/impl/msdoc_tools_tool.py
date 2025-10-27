"""
Copyright All rights Reserved 2025-2030, Ashutosh Sinha, Email: ajsinha@gmail.com
Microsoft Documents (Word & Excel) MCP Tool Implementation
"""

import json
import os
import re
from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
from tools.base_mcp_tool import BaseMCPTool


class MsDocToolsTool(BaseMCPTool):
    """
    Microsoft Documents tool for Word (.docx) and Excel (.xlsx) file operations
    """
    
    def __init__(self, config: Dict = None):
        """Initialize MS Doc Tools"""
        default_config = {
            'name': 'msdoc_tools',
            'description': 'Work with Microsoft Word and Excel documents',
            'version': '1.0.0',
            'enabled': True
        }
        if config:
            default_config.update(config)
        super().__init__(default_config)
        
        # Documents directory
        self.docs_dir = config.get('docs_directory', 'data/msdocs') if config else 'data/msdocs'
        
        # Ensure directory exists
        os.makedirs(self.docs_dir, exist_ok=True)
    
    def get_input_schema(self) -> Dict:
        """Get input schema for MS Doc Tools"""
        return {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "description": "Action to perform",
                    "enum": [
                        "list_files",
                        "read_word",
                        "read_excel",
                        "search_word",
                        "search_excel",
                        "get_word_metadata",
                        "get_excel_metadata",
                        "extract_text",
                        "get_excel_sheets",
                        "read_excel_sheet"
                    ]
                },
                "filename": {
                    "type": "string",
                    "description": "Name of the file to operate on"
                },
                "file_type": {
                    "type": "string",
                    "description": "Type of files to list",
                    "enum": ["all", "word", "excel"]
                },
                "search_term": {
                    "type": "string",
                    "description": "Search term to find in documents"
                },
                "sheet_name": {
                    "type": "string",
                    "description": "Name of Excel sheet to read"
                },
                "sheet_index": {
                    "type": "integer",
                    "description": "Index of Excel sheet to read (0-based)",
                    "minimum": 0
                },
                "max_rows": {
                    "type": "integer",
                    "description": "Maximum number of rows to return from Excel",
                    "default": 100,
                    "minimum": 1,
                    "maximum": 10000
                },
                "include_formulas": {
                    "type": "boolean",
                    "description": "Include cell formulas in Excel data",
                    "default": False
                }
            },
            "required": ["action"]
        }
    
    def execute(self, arguments: Dict[str, Any]) -> Any:
        """
        Execute MS Doc Tools action
        
        Args:
            arguments: Tool arguments
            
        Returns:
            Result based on action
        """
        action = arguments.get('action')
        
        if action == 'list_files':
            file_type = arguments.get('file_type', 'all')
            return self._list_files(file_type)
            
        elif action == 'read_word':
            filename = arguments.get('filename')
            if not filename:
                raise ValueError("'filename' is required for read_word")
            return self._read_word_document(filename)
            
        elif action == 'read_excel':
            filename = arguments.get('filename')
            if not filename:
                raise ValueError("'filename' is required for read_excel")
            max_rows = arguments.get('max_rows', 100)
            include_formulas = arguments.get('include_formulas', False)
            return self._read_excel_file(filename, max_rows, include_formulas)
            
        elif action == 'search_word':
            filename = arguments.get('filename')
            search_term = arguments.get('search_term')
            if not filename or not search_term:
                raise ValueError("'filename' and 'search_term' are required")
            return self._search_word_document(filename, search_term)
            
        elif action == 'search_excel':
            filename = arguments.get('filename')
            search_term = arguments.get('search_term')
            if not filename or not search_term:
                raise ValueError("'filename' and 'search_term' are required")
            return self._search_excel_file(filename, search_term)
            
        elif action == 'get_word_metadata':
            filename = arguments.get('filename')
            if not filename:
                raise ValueError("'filename' is required")
            return self._get_word_metadata(filename)
            
        elif action == 'get_excel_metadata':
            filename = arguments.get('filename')
            if not filename:
                raise ValueError("'filename' is required")
            return self._get_excel_metadata(filename)
            
        elif action == 'extract_text':
            filename = arguments.get('filename')
            if not filename:
                raise ValueError("'filename' is required")
            return self._extract_text(filename)
            
        elif action == 'get_excel_sheets':
            filename = arguments.get('filename')
            if not filename:
                raise ValueError("'filename' is required")
            return self._get_excel_sheets(filename)
            
        elif action == 'read_excel_sheet':
            filename = arguments.get('filename')
            sheet_name = arguments.get('sheet_name')
            sheet_index = arguments.get('sheet_index')
            max_rows = arguments.get('max_rows', 100)
            
            if not filename:
                raise ValueError("'filename' is required")
            if sheet_name is None and sheet_index is None:
                raise ValueError("Either 'sheet_name' or 'sheet_index' is required")
                
            return self._read_excel_sheet(filename, sheet_name, sheet_index, max_rows)
            
        else:
            raise ValueError(f"Unknown action: {action}")
    
    def _list_files(self, file_type: str = 'all') -> Dict:
        """
        List files in the documents directory
        
        Args:
            file_type: Type of files to list
            
        Returns:
            List of files with metadata
        """
        files = []
        path = Path(self.docs_dir)
        
        # Define extensions based on file type
        if file_type == 'word':
            extensions = ['.docx']
        elif file_type == 'excel':
            extensions = ['.xlsx', '.xlsm']
        else:  # all
            extensions = ['.docx', '.xlsx', '.xlsm']
        
        for ext in extensions:
            for file_path in path.glob(f'*{ext}'):
                if file_path.is_file():
                    stat = file_path.stat()
                    files.append({
                        'filename': file_path.name,
                        'type': 'word' if ext == '.docx' else 'excel',
                        'size': stat.st_size,
                        'size_mb': round(stat.st_size / (1024 * 1024), 2),
                        'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        'path': str(file_path)
                    })
        
        # Sort by modified date (newest first)
        files.sort(key=lambda x: x['modified'], reverse=True)
        
        return {
            'directory': self.docs_dir,
            'file_type': file_type,
            'count': len(files),
            'files': files
        }
    
    def _read_word_document(self, filename: str) -> Dict:
        """
        Read Word document content
        
        Args:
            filename: Name of Word file
            
        Returns:
            Document content and metadata
        """
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        if not filename.endswith('.docx'):
            raise ValueError("File must be a .docx file")
        
        try:
            # Try to import python-docx
            try:
                from docx import Document
            except ImportError:
                return self._read_word_basic(file_path)
            
            doc = Document(str(file_path))
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Get core properties
            props = doc.core_properties
            
            return {
                'filename': filename,
                'type': 'word',
                'paragraph_count': len(paragraphs),
                'table_count': len(tables),
                'paragraphs': paragraphs,
                'tables': tables,
                'metadata': {
                    'author': props.author if props.author else 'Unknown',
                    'title': props.title if props.title else '',
                    'subject': props.subject if props.subject else '',
                    'created': props.created.isoformat() if props.created else None,
                    'modified': props.modified.isoformat() if props.modified else None
                }
            }
            
        except Exception as e:
            self.logger.error(f"Error reading Word document: {e}")
            raise ValueError(f"Failed to read Word document: {str(e)}")
    
    def _read_word_basic(self, file_path: Path) -> Dict:
        """Basic Word document reading without python-docx"""
        return {
            'filename': file_path.name,
            'type': 'word',
            'error': 'python-docx library not available',
            'note': 'Install python-docx for full functionality',
            'file_exists': True,
            'size': file_path.stat().st_size
        }
    
    def _read_excel_file(
        self,
        filename: str,
        max_rows: int = 100,
        include_formulas: bool = False
    ) -> Dict:
        """
        Read Excel file content
        
        Args:
            filename: Name of Excel file
            max_rows: Maximum rows to return
            include_formulas: Include cell formulas
            
        Returns:
            Excel data and metadata
        """
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        if not (filename.endswith('.xlsx') or filename.endswith('.xlsm')):
            raise ValueError("File must be an .xlsx or .xlsm file")
        
        try:
            # Try to import openpyxl
            try:
                from openpyxl import load_workbook
            except ImportError:
                return self._read_excel_basic(file_path)
            
            wb = load_workbook(str(file_path), data_only=not include_formulas)
            
            # Get all sheets
            sheets_data = {}
            
            for sheet_name in wb.sheetnames[:5]:  # Limit to first 5 sheets
                sheet = wb[sheet_name]
                
                # Get data
                data = []
                row_count = 0
                
                for row in sheet.iter_rows(values_only=not include_formulas):
                    if row_count >= max_rows:
                        break
                    
                    # Convert row to list, handling None values
                    row_data = [cell if cell is not None else '' for cell in row]
                    data.append(row_data)
                    row_count += 1
                
                sheets_data[sheet_name] = {
                    'rows': data,
                    'row_count': row_count,
                    'column_count': sheet.max_column,
                    'total_rows': sheet.max_row
                }
            
            return {
                'filename': filename,
                'type': 'excel',
                'sheet_count': len(wb.sheetnames),
                'sheet_names': wb.sheetnames,
                'sheets': sheets_data,
                'note': f'Showing up to {max_rows} rows per sheet'
            }
            
        except Exception as e:
            self.logger.error(f"Error reading Excel file: {e}")
            raise ValueError(f"Failed to read Excel file: {str(e)}")
    
    def _read_excel_basic(self, file_path: Path) -> Dict:
        """Basic Excel file reading without openpyxl"""
        return {
            'filename': file_path.name,
            'type': 'excel',
            'error': 'openpyxl library not available',
            'note': 'Install openpyxl for full functionality',
            'file_exists': True,
            'size': file_path.stat().st_size
        }
    
    def _search_word_document(self, filename: str, search_term: str) -> Dict:
        """
        Search for text in Word document
        
        Args:
            filename: Name of Word file
            search_term: Text to search for
            
        Returns:
            Search results with matches
        """
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from docx import Document
        except ImportError:
            return {
                'error': 'python-docx library not available',
                'note': 'Install python-docx to search Word documents'
            }
        
        try:
            doc = Document(str(file_path))
            
            matches = []
            search_pattern = re.compile(re.escape(search_term), re.IGNORECASE)
            
            # Search in paragraphs
            for i, para in enumerate(doc.paragraphs):
                if search_pattern.search(para.text):
                    # Get context (50 chars before and after)
                    text = para.text
                    for match in search_pattern.finditer(text):
                        start = max(0, match.start() - 50)
                        end = min(len(text), match.end() + 50)
                        context = text[start:end]
                        
                        matches.append({
                            'type': 'paragraph',
                            'index': i,
                            'context': context,
                            'position': match.start()
                        })
            
            # Search in tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if search_pattern.search(cell.text):
                            matches.append({
                                'type': 'table',
                                'table_index': table_idx,
                                'row': row_idx,
                                'column': cell_idx,
                                'text': cell.text[:100]
                            })
            
            return {
                'filename': filename,
                'search_term': search_term,
                'match_count': len(matches),
                'matches': matches
            }
            
        except Exception as e:
            self.logger.error(f"Error searching Word document: {e}")
            raise ValueError(f"Failed to search document: {str(e)}")
    
    def _search_excel_file(self, filename: str, search_term: str) -> Dict:
        """
        Search for text in Excel file
        
        Args:
            filename: Name of Excel file
            search_term: Text to search for
            
        Returns:
            Search results with matches
        """
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from openpyxl import load_workbook
        except ImportError:
            return {
                'error': 'openpyxl library not available',
                'note': 'Install openpyxl to search Excel files'
            }
        
        try:
            wb = load_workbook(str(file_path), data_only=True)
            
            matches = []
            search_lower = search_term.lower()
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                for row_idx, row in enumerate(sheet.iter_rows(), 1):
                    for col_idx, cell in enumerate(row, 1):
                        if cell.value and search_lower in str(cell.value).lower():
                            matches.append({
                                'sheet': sheet_name,
                                'row': row_idx,
                                'column': col_idx,
                                'cell': f'{cell.column_letter}{cell.row}',
                                'value': str(cell.value)
                            })
            
            return {
                'filename': filename,
                'search_term': search_term,
                'match_count': len(matches),
                'matches': matches[:100]  # Limit to 100 matches
            }
            
        except Exception as e:
            self.logger.error(f"Error searching Excel file: {e}")
            raise ValueError(f"Failed to search file: {str(e)}")
    
    def _get_word_metadata(self, filename: str) -> Dict:
        """Get Word document metadata"""
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from docx import Document
        except ImportError:
            return self._get_basic_metadata(file_path)
        
        try:
            doc = Document(str(file_path))
            props = doc.core_properties
            
            return {
                'filename': filename,
                'type': 'word',
                'metadata': {
                    'author': props.author if props.author else 'Unknown',
                    'title': props.title if props.title else '',
                    'subject': props.subject if props.subject else '',
                    'keywords': props.keywords if props.keywords else '',
                    'created': props.created.isoformat() if props.created else None,
                    'modified': props.modified.isoformat() if props.modified else None,
                    'last_modified_by': props.last_modified_by if props.last_modified_by else ''
                },
                'statistics': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'sections': len(doc.sections)
                }
            }
            
        except Exception as e:
            return self._get_basic_metadata(file_path)
    
    def _get_excel_metadata(self, filename: str) -> Dict:
        """Get Excel file metadata"""
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from openpyxl import load_workbook
        except ImportError:
            return self._get_basic_metadata(file_path)
        
        try:
            wb = load_workbook(str(file_path), data_only=True)
            props = wb.properties
            
            sheet_info = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_info.append({
                    'name': sheet_name,
                    'rows': sheet.max_row,
                    'columns': sheet.max_column
                })
            
            return {
                'filename': filename,
                'type': 'excel',
                'metadata': {
                    'creator': props.creator if props.creator else 'Unknown',
                    'title': props.title if props.title else '',
                    'subject': props.subject if props.subject else '',
                    'created': props.created.isoformat() if props.created else None,
                    'modified': props.modified.isoformat() if props.modified else None
                },
                'statistics': {
                    'sheet_count': len(wb.sheetnames),
                    'sheets': sheet_info
                }
            }
            
        except Exception as e:
            return self._get_basic_metadata(file_path)
    
    def _get_basic_metadata(self, file_path: Path) -> Dict:
        """Get basic file metadata"""
        stat = file_path.stat()
        return {
            'filename': file_path.name,
            'size': stat.st_size,
            'size_mb': round(stat.st_size / (1024 * 1024), 2),
            'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'note': 'Limited metadata - install python-docx or openpyxl for full details'
        }
    
    def _extract_text(self, filename: str) -> Dict:
        """
        Extract all text from document
        
        Args:
            filename: Name of file
            
        Returns:
            Extracted text
        """
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        if filename.endswith('.docx'):
            return self._extract_text_word(file_path)
        elif filename.endswith('.xlsx') or filename.endswith('.xlsm'):
            return self._extract_text_excel(file_path)
        else:
            raise ValueError("Unsupported file type")
    
    def _extract_text_word(self, file_path: Path) -> Dict:
        """Extract text from Word document"""
        try:
            from docx import Document
        except ImportError:
            return {'error': 'python-docx not available'}
        
        doc = Document(str(file_path))
        
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = '\n'.join(paragraphs)
        
        return {
            'filename': file_path.name,
            'type': 'word',
            'text': full_text,
            'character_count': len(full_text),
            'paragraph_count': len(paragraphs),
            'word_count': len(full_text.split())
        }
    
    def _extract_text_excel(self, file_path: Path) -> Dict:
        """Extract text from Excel file"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            return {'error': 'openpyxl not available'}
        
        wb = load_workbook(str(file_path), data_only=True)
        
        all_text = []
        cell_count = 0
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        all_text.append(str(cell.value))
                        cell_count += 1
        
        full_text = '\n'.join(all_text)
        
        return {
            'filename': file_path.name,
            'type': 'excel',
            'text': full_text,
            'character_count': len(full_text),
            'cell_count': cell_count,
            'sheet_count': len(wb.sheetnames)
        }
    
    def _get_excel_sheets(self, filename: str) -> Dict:
        """Get list of sheets in Excel file"""
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from openpyxl import load_workbook
        except ImportError:
            return {'error': 'openpyxl not available'}
        
        wb = load_workbook(str(file_path), data_only=True)
        
        sheets = []
        for idx, sheet_name in enumerate(wb.sheetnames):
            sheet = wb[sheet_name]
            sheets.append({
                'index': idx,
                'name': sheet_name,
                'rows': sheet.max_row,
                'columns': sheet.max_column
            })
        
        return {
            'filename': filename,
            'sheet_count': len(sheets),
            'sheets': sheets
        }
    
    def _read_excel_sheet(
        self,
        filename: str,
        sheet_name: Optional[str],
        sheet_index: Optional[int],
        max_rows: int
    ) -> Dict:
        """Read specific Excel sheet"""
        file_path = Path(self.docs_dir) / filename
        
        if not file_path.exists():
            raise ValueError(f"File not found: {filename}")
        
        try:
            from openpyxl import load_workbook
        except ImportError:
            return {'error': 'openpyxl not available'}
        
        wb = load_workbook(str(file_path), data_only=True)
        
        # Get sheet by name or index
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                raise ValueError(f"Sheet '{sheet_name}' not found")
            sheet = wb[sheet_name]
        else:
            if sheet_index >= len(wb.sheetnames):
                raise ValueError(f"Sheet index {sheet_index} out of range")
            sheet_name = wb.sheetnames[sheet_index]
            sheet = wb[sheet_name]
        
        # Get data
        data = []
        for idx, row in enumerate(sheet.iter_rows(values_only=True)):
            if idx >= max_rows:
                break
            row_data = [cell if cell is not None else '' for cell in row]
            data.append(row_data)
        
        return {
            'filename': filename,
            'sheet_name': sheet_name,
            'sheet_index': wb.sheetnames.index(sheet_name),
            'rows': data,
            'row_count': len(data),
            'total_rows': sheet.max_row,
            'column_count': sheet.max_column
        }
