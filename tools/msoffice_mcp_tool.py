"""
MS Office MCP Tool implementation for Word and Excel files
"""
import os
from typing import Dict, Any, List
from .base_mcp_tool import BaseMCPTool
import docx
import openpyxl
import pandas as pd
from datetime import datetime


class MSOfficeMCPTool(BaseMCPTool):
    """MCP Tool for MS Office file operations"""

    def _initialize(self):
        """Initialize MS Office specific components"""
        self.data_folder = os.environ.get('OFFICE_DATA_FOLDER', 'data/office_files')

        # Create folder if it doesn't exist
        os.makedirs(self.data_folder, exist_ok=True)

        # Cache for loaded documents
        self.document_cache = {}
        self.excel_cache = {}

        # Scan for available files
        self._scan_files()

    def _scan_files(self):
        """Scan the data folder for Office files"""
        self.available_files = {
            'word': [],
            'excel': []
        }

        if os.path.exists(self.data_folder):
            for file in os.listdir(self.data_folder):
                file_path = os.path.join(self.data_folder, file)
                if file.endswith('.docx'):
                    self.available_files['word'].append(file)
                elif file.endswith('.xlsx') or file.endswith('.xls'):
                    self.available_files['excel'].append(file)

    def handle_tool_call(self, tool_name: str, arguments: Dict[str, Any]) -> Dict[str, Any]:
        """Handle MS Office tool calls"""
        try:
            if self.check_rate_limit():
                error_msg = "Rate limit exceeded"
                self.record_call(tool_name, arguments, error=error_msg)
                return {"error": error_msg, "status": 429}

            result = None

            tool_methods = {
                "list_files": self._list_files,
                "read_word_document": self._read_word_document,
                "search_word_documents": self._search_word_documents,
                "get_word_metadata": self._get_word_metadata,
                "read_excel_sheet": self._read_excel_sheet,
                "get_excel_sheets": self._get_excel_sheets,
                "query_excel_data": self._query_excel_data,
                "get_excel_statistics": self._get_excel_statistics,
                "search_excel_files": self._search_excel_files,
                "create_word_document": self._create_word_document,
                "create_excel_file": self._create_excel_file
            }

            if tool_name in tool_methods:
                result = tool_methods[tool_name](arguments)
            else:
                result = {"error": f"Unknown tool: {tool_name}"}

            self.record_call(tool_name, arguments, result=result)
            return result

        except Exception as e:
            error_msg = str(e)
            self.record_call(tool_name, arguments, error=error_msg)
            return {"error": error_msg, "status": 500}

    def _list_files(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """List available Office files"""
        file_type = params.get('file_type', 'all')  # 'word', 'excel', 'all'

        # Rescan files
        self._scan_files()

        if file_type == 'word':
            files = self.available_files['word']
        elif file_type == 'excel':
            files = self.available_files['excel']
        else:
            files = self.available_files['word'] + self.available_files['excel']

        # Get file details
        file_details = []
        for filename in files:
            file_path = os.path.join(self.data_folder, filename)
            if os.path.exists(file_path):
                stats = os.stat(file_path)
                file_details.append({
                    'name': filename,
                    'size_bytes': stats.st_size,
                    'modified': datetime.fromtimestamp(stats.st_mtime).isoformat(),
                    'type': 'Word' if filename.endswith('.docx') else 'Excel'
                })

        return {
            "files": file_details,
            "count": len(file_details),
            "folder": self.data_folder
        }

    def _read_word_document(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Read content from a Word document"""
        filename = params.get('filename', '')
        max_paragraphs = params.get('max_paragraphs', 50)

        if not filename:
            return {"error": "Filename is required"}

        file_path = os.path.join(self.data_folder, filename)

        if not os.path.exists(file_path):
            return {"error": f"File '{filename}' not found"}

        try:
            # Load document (use cache if available)
            if filename not in self.document_cache:
                self.document_cache[filename] = docx.Document(file_path)

            doc = self.document_cache[filename]

            # Extract content
            content = {
                'paragraphs': [],
                'tables': [],
                'headers': [],
                'footers': []
            }

            # Get paragraphs
            for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
                if para.text.strip():
                    content['paragraphs'].append({
                        'index': i,
                        'text': para.text,
                        'style': para.style.name if para.style else None
                    })

            # Get tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append({
                    'index': i,
                    'rows': len(table.rows),
                    'columns': len(table.columns),
                    'data': table_data[:10]  # Limit to first 10 rows
                })

            # Get headers and footers
            for section in doc.sections:
                if section.header:
                    header_text = '\n'.join([p.text for p in section.header.paragraphs])
                    if header_text.strip():
                        content['headers'].append(header_text)

                if section.footer:
                    footer_text = '\n'.join([p.text for p in section.footer.paragraphs])
                    if footer_text.strip():
                        content['footers'].append(footer_text)

            return {
                "filename": filename,
                "content": content,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }

        except Exception as e:
            return {"error": str(e)}

    def _search_word_documents(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Search for text in Word documents"""
        search_term = params.get('search_term', '')
        case_sensitive = params.get('case_sensitive', False)

        if not search_term:
            return {"error": "Search term is required"}

        results = []

        for filename in self.available_files['word']:
            file_path = os.path.join(self.data_folder, filename)

            try:
                doc = docx.Document(file_path)
                matches = []

                # Search in paragraphs
                for i, para in enumerate(doc.paragraphs):
                    text = para.text
                    search_text = search_term

                    if not case_sensitive:
                        text = text.lower()
                        search_text = search_text.lower()

                    if search_text in text:
                        matches.append({
                            'type': 'paragraph',
                            'index': i,
                            'text': para.text[:200]  # First 200 chars
                        })

                # Search in tables
                for i, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        for col_idx, cell in enumerate(row.cells):
                            text = cell.text
                            search_text = search_term

                            if not case_sensitive:
                                text = text.lower()
                                search_text = search_text.lower()

                            if search_text in text:
                                matches.append({
                                    'type': 'table',
                                    'table_index': i,
                                    'row': row_idx,
                                    'column': col_idx,
                                    'text': cell.text[:200]
                                })

                if matches:
                    results.append({
                        'filename': filename,
                        'matches': matches[:10],  # Limit to 10 matches per file
                        'total_matches': len(matches)
                    })

            except Exception as e:
                results.append({
                    'filename': filename,
                    'error': str(e)
                })

        return {
            "search_term": search_term,
            "results": results,
            "files_searched": len(self.available_files['word']),
            "files_with_matches": len([r for r in results if 'matches' in r])
        }

    def _get_word_metadata(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Get metadata from a Word document"""
        filename = params.get('filename', '')

        if not filename:
            return {"error": "Filename is required"}

        file_path = os.path.join(self.data_folder, filename)

        if not os.path.exists(file_path):
            return {"error": f"File '{filename}' not found"}

        try:
            doc = docx.Document(file_path)
            core_properties = doc.core_properties

            metadata = {
                'title': core_properties.title,
                'author': core_properties.author,
                'subject': core_properties.subject,
                'keywords': core_properties.keywords,
                'created': core_properties.created.isoformat() if core_properties.created else None,
                'modified': core_properties.modified.isoformat() if core_properties.modified else None,
                'last_modified_by': core_properties.last_modified_by,
                'revision': core_properties.revision,
                'category': core_properties.category,
                'comments': core_properties.comments
            }

            # Document statistics
            stats = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
                'word_count': sum(len(p.text.split()) for p in doc.paragraphs)
            }

            return {
                "filename": filename,
                "metadata": metadata,
                "statistics": stats
            }

        except Exception as e:
            return {"error": str(e)}

    def _read_excel_sheet(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Read data from an Excel sheet"""
        filename = params.get('filename', '')
        sheet_name = params.get('sheet_name', None)
        max_rows = params.get('max_rows', 100)

        if not filename:
            return {"error": "Filename is required"}

        file_path = os.path.join(self.data_folder, filename)

        if not os.path.exists(file_path):
            return {"error": f"File '{filename}' not found"}

        try:
            # Load Excel file
            if filename not in self.excel_cache:
                self.excel_cache[filename] = pd.ExcelFile(file_path)

            excel_file = self.excel_cache[filename]

            # Get sheet name
            if sheet_name is None:
                sheet_name = excel_file.sheet_names[0]
            elif sheet_name not in excel_file.sheet_names:
                return {"error": f"Sheet '{sheet_name}' not found"}

            # Read data
            df = pd.read_excel(excel_file, sheet_name=sheet_name, nrows=max_rows)

            # Convert to dictionary
            data = {
                'columns': df.columns.tolist(),
                'data': df.head(max_rows).to_dict('records'),
                'shape': {
                    'rows': len(df),
                    'columns': len(df.columns)
                },
                'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()}
            }

            return {
                "filename": filename,
                "sheet_name": sheet_name,
                "data": data
            }

        except Exception as e:
            return {"error": str(e)}

    def _get_excel_sheets(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Get list of sheets in an Excel file"""
        filename = params.get('filename', '')

        if not filename:
            return {"error": "Filename is required"}

        file_path = os.path.join(self.data_folder, filename)

        if not os.path.exists(file_path):
            return {"error": f"File '{filename}' not found"}

        try:
            wb = openpyxl.load_workbook(file_path, read_only=True)

            sheets = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheets.append({
                    'name': sheet_name,
                    'max_row': sheet.max_row,
                    'max_column': sheet.max_column
                })

            wb.close()

            return {
                "filename": filename,
                "sheets": sheets,
                "sheet_count": len(sheets)
            }

        except Exception as e:
            return {"error": str(e)}

    def _query_excel_data(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Query Excel data with filters"""
        filename = params.get('filename', '')
        sheet_name = params.get('sheet_name', None)
        filters = params.get('filters', {})  # Column-value pairs
        columns = params.get('columns', [])  # Columns to return

        if not filename:
            return {"error": "Filename is required"}

        file_path = os.path.join(self.data_folder, filename)

        if not os.path.exists(file_path):
            return {"error": f"File '{filename}' not found"}

        try:
            # Read Excel data
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # Apply filters
            for column, value in filters.items():
                if column in df.columns:
                    df = df[df[column] == value]

            # Select columns
            if columns and all(col in df.columns for col in columns):
                df = df[columns]

            # Convert to dictionary
            results = {
                'data': df.to_dict('records'),
                'row_count': len(df),
                'columns': df.columns.tolist()
            }

            return {
                "filename": filename,
                "sheet_name": sheet_name,
                "filters": filters,
                "results": results
            }

        except Exception as e:
            return {"error": str(e)}

    def _get_excel_statistics(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Get statistics from Excel data"""
        filename = params.get('filename', '')
        sheet_name = params.get('sheet_name', None)

        if not filename:
            return {"error": "Filename is required"}

        file_path = os.path.join(self.data_folder, filename)

        if not os.path.exists(file_path):
            return {"error": f"File '{filename}' not found"}

        try:
            # Read Excel data
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # Get statistics
            stats = {
                'shape': {
                    'rows': len(df),
                    'columns': len(df.columns)
                },
                'columns': df.columns.tolist(),
                'numeric_columns': df.select_dtypes(include=['number']).columns.tolist(),
                'null_counts': df.isnull().sum().to_dict(),
                'unique_counts': {col: df[col].nunique() for col in df.columns}
            }

            # Numeric statistics
            numeric_stats = {}
            for col in df.select_dtypes(include=['number']).columns:
                numeric_stats[col] = {
                    'mean': float(df[col].mean()) if not df[col].isna().all() else None,
                    'median': float(df[col].median()) if not df[col].isna().all() else None,
                    'min': float(df[col].min()) if not df[col].isna().all() else None,
                    'max': float(df[col].max()) if not df[col].isna().all() else None,
                    'std': float(df[col].std()) if not df[col].isna().all() else None
                }

            stats['numeric_statistics'] = numeric_stats

            return {
                "filename": filename,
                "sheet_name": sheet_name,
                "statistics": stats
            }

        except Exception as e:
            return {"error": str(e)}

    def _search_excel_files(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Search for values in Excel files"""
        search_value = params.get('search_value', '')

        if not search_value:
            return {"error": "Search value is required"}

        results = []

        for filename in self.available_files['excel']:
            file_path = os.path.join(self.data_folder, filename)

            try:
                excel_file = pd.ExcelFile(file_path)
                file_results = []

                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)

                    # Search for the value in all columns
                    matches = []
                    for col in df.columns:
                        mask = df[col].astype(str).str.contains(str(search_value), case=False, na=False)
                        if mask.any():
                            matching_indices = df[mask].index.tolist()
                            matches.append({
                                'column': col,
                                'row_indices': matching_indices[:10],  # Limit to 10
                                'match_count': mask.sum()
                            })

                    if matches:
                        file_results.append({
                            'sheet': sheet_name,
                            'matches': matches
                        })

                if file_results:
                    results.append({
                        'filename': filename,
                        'results': file_results
                    })

            except Exception as e:
                results.append({
                    'filename': filename,
                    'error': str(e)
                })

        return {
            "search_value": search_value,
            "results": results,
            "files_searched": len(self.available_files['excel'])
        }

    def _create_word_document(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Create a new Word document"""
        filename = params.get('filename', 'new_document.docx')
        content = params.get('content', [])  # List of paragraphs

        if not filename.endswith('.docx'):
            filename += '.docx'

        file_path = os.path.join(self.data_folder, filename)

        try:
            doc = docx.Document()

            for item in content:
                if isinstance(item, str):
                    doc.add_paragraph(item)
                elif isinstance(item, dict):
                    if item.get('type') == 'heading':
                        doc.add_heading(item.get('text', ''), level=item.get('level', 1))
                    elif item.get('type') == 'paragraph':
                        doc.add_paragraph(item.get('text', ''))

            doc.save(file_path)

            # Update file list
            self._scan_files()

            return {
                "filename": filename,
                "path": file_path,
                "message": "Document created successfully"
            }

        except Exception as e:
            return {"error": str(e)}

    def _create_excel_file(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Create a new Excel file"""
        filename = params.get('filename', 'new_spreadsheet.xlsx')
        data = params.get('data', {})  # Dict of sheet_name: data

        if not filename.endswith(('.xlsx', '.xls')):
            filename += '.xlsx'

        file_path = os.path.join(self.data_folder, filename)

        try:
            with pd.ExcelWriter(file_path) as writer:
                for sheet_name, sheet_data in data.items():
                    df = pd.DataFrame(sheet_data)
                    df.to_excel(writer, sheet_name=sheet_name, index=False)

            # Update file list
            self._scan_files()

            return {
                "filename": filename,
                "path": file_path,
                "sheets": list(data.keys()),
                "message": "Excel file created successfully"
            }

        except Exception as e:
            return {"error": str(e)}